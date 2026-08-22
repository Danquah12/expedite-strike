"""
PT Report Export Callbacks — CSV / PDF / Word
Inputs:  pt-download-csv-btn, pt-download-pdf-btn, pt-download-word-btn
Outputs: pt-dl-csv, pt-dl-pdf, pt-dl-word  (dcc.Download sinks in layout_post_exploitation)
"""
import base64, csv, io, datetime as _dt
import dash
from dash import callback, Output, Input
from dash.exceptions import PreventUpdate


def _fetch_report_rows():
    """Query Neo4j and return list-of-dicts for export."""
    try:
        from cyber_range.services.neo4j_engine import Neo4jEngine
        ng = Neo4jEngine()
        with ng.driver.session() as s:
            rows = s.run("""
                MATCH (h:Host)
                WHERE h.ip IS NOT NULL
                OPTIONAL MATCH (h)-[:RUNS_SERVICE|EXPOSES|HasService]->(svc:Service)
                OPTIONAL MATCH (svc)-[:HAS_FINDING|HAS_VULN]->(f:Finding)
                RETURN
                  coalesce(h.host, h.ip)              AS host,
                  h.ip                                AS ip,
                  coalesce(svc.name,'-')              AS service,
                  coalesce(toString(svc.port),'-')    AS port,
                  coalesce(f.severity,'-')            AS severity,
                  coalesce(f.name,'-')                AS finding,
                  coalesce(f.cve,'-')                 AS cve,
                  coalesce(toString(f.cvss),'-')      AS cvss,
                  coalesce(f.scanner,'-')             AS scanner
                ORDER BY host, severity
            """).data()
        ng.driver.close()
        return rows
    except Exception as e:
        print(f"[pt-report] Neo4j fetch error: {e}")
        return []


# ── CSV ───────────────────────────────────────────────────────────────────────
@callback(
    Output("pt-dl-csv",        "data"),
    Output("pt-report-status", "children", allow_duplicate=True),
    Input("pt-download-csv-btn", "n_clicks"),
    prevent_initial_call=True,
)
def download_pt_csv(n):
    if not n:
        raise PreventUpdate
    rows = _fetch_report_rows()
    if not rows:
        return dash.no_update, "No data returned from Neo4j."
    buf = io.StringIO()
    writer = csv.DictWriter(buf, fieldnames=list(rows[0].keys()))
    writer.writeheader()
    writer.writerows(rows)
    ts = _dt.datetime.now().strftime("%Y%m%d_%H%M%S")
    return (dict(content=buf.getvalue(),
                 filename=f"pentest_report_{ts}.csv",
                 type="text/csv"),
            f"CSV ready — {len(rows)} rows exported.")


# ── PDF ───────────────────────────────────────────────────────────────────────
@callback(
    Output("pt-dl-pdf",        "data"),
    Output("pt-report-status", "children", allow_duplicate=True),
    Input("pt-download-pdf-btn", "n_clicks"),
    prevent_initial_call=True,
)
def download_pt_pdf(n):
    if not n:
        raise PreventUpdate
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib import colors
        from reportlab.platypus import (SimpleDocTemplate, Table, TableStyle,
                                        Paragraph, Spacer)
        from reportlab.lib.styles import getSampleStyleSheet
    except ImportError:
        return dash.no_update, "reportlab not installed — run: pip install reportlab"

    rows = _fetch_report_rows()
    if not rows:
        return dash.no_update, "No data returned from Neo4j."

    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4,
                            leftMargin=30, rightMargin=30,
                            topMargin=40, bottomMargin=30)
    styles = getSampleStyleSheet()
    ts = _dt.datetime.now().strftime("%Y-%m-%d %H:%M:%S")

    sev_cnts = {}
    for r in rows:
        sv = r.get("severity", "-")
        sev_cnts[sv] = sev_cnts.get(sv, 0) + 1

    story = [
        Paragraph("PENETRATION TEST REPORT", styles["Title"]),
        Paragraph(f"Generated: {ts}  |  Source: Live Neo4j", styles["Normal"]),
        Spacer(1, 12),
        Paragraph("Severity Summary", styles["Heading2"]),
    ]
    for sv, cnt in sev_cnts.items():
        story.append(Paragraph(f"  - {sv}: {cnt}", styles["Normal"]))
    story.append(Spacer(1, 10))
    story.append(Paragraph("Findings Detail", styles["Heading2"]))

    headers = ["Host/IP", "Service", "Port", "Severity", "Finding", "CVE", "CVSS", "Scanner"]
    tbl_data = [headers] + [
        [r.get("host",""), r.get("service",""), r.get("port",""),
         r.get("severity",""), (r.get("finding","") or "")[:50],
         r.get("cve",""), r.get("cvss",""), r.get("scanner","")]
        for r in rows
    ]
    col_w = [95, 55, 35, 55, 145, 70, 35, 55]
    tbl = Table(tbl_data, colWidths=col_w, repeatRows=1)
    tbl.setStyle(TableStyle([
        ("BACKGROUND",     (0, 0), (-1, 0), colors.HexColor("#1a237e")),
        ("TEXTCOLOR",      (0, 0), (-1, 0), colors.white),
        ("FONTSIZE",       (0, 0), (-1, -1), 7),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.HexColor("#1c1c1c"),
                                               colors.HexColor("#141414")]),
        ("TEXTCOLOR",      (0, 1), (-1, -1), colors.HexColor("#cccccc")),
        ("GRID",           (0, 0), (-1, -1), 0.3, colors.HexColor("#333")),
        ("LEFTPADDING",    (0, 0), (-1, -1), 4),
        ("RIGHTPADDING",   (0, 0), (-1, -1), 4),
    ]))
    story.append(tbl)
    doc.build(story)

    b64 = base64.b64encode(buf.getvalue()).decode()
    ts2 = _dt.datetime.now().strftime("%Y%m%d_%H%M%S")
    return (dict(content=b64, filename=f"pentest_report_{ts2}.pdf",
                 type="application/pdf", base64=True),
            f"PDF ready — {len(rows)} findings exported.")


# ── Word (.docx) ──────────────────────────────────────────────────────────────
@callback(
    Output("pt-dl-word",        "data"),
    Output("pt-report-status",  "children", allow_duplicate=True),
    Input("pt-download-word-btn", "n_clicks"),
    prevent_initial_call=True,
)
def download_pt_word(n):
    if not n:
        raise PreventUpdate
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
    except ImportError:
        return dash.no_update, "python-docx not installed — run: pip install python-docx"

    rows = _fetch_report_rows()
    if not rows:
        return dash.no_update, "No data returned from Neo4j."

    doc = Document()
    ts = _dt.datetime.now().strftime("%Y-%m-%d %H:%M:%S")

    ttl = doc.add_heading("Penetration Test Report", 0)
    ttl.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"Generated: {ts}  |  Source: Live Neo4j")
    doc.add_paragraph(f"Total Findings: {len(rows)}")

    doc.add_heading("Executive Summary", level=1)
    sev_cnts = {}
    for r in rows:
        sv = r.get("severity", "Unknown")
        sev_cnts[sv] = sev_cnts.get(sv, 0) + 1
    for sv in ["Critical", "High", "Medium", "Low", "Info", "-"]:
        if sv in sev_cnts:
            doc.add_paragraph(f"{sv}: {sev_cnts[sv]} finding(s)", style="List Bullet")

    doc.add_heading("Detailed Findings", level=1)
    hdrs = ["Host/IP", "Service", "Port", "Severity", "Finding", "CVE", "CVSS", "Scanner"]
    tbl = doc.add_table(rows=1, cols=len(hdrs))
    tbl.style = "Table Grid"
    hrow = tbl.rows[0]
    for i, h in enumerate(hdrs):
        c = hrow.cells[i]
        c.text = h
        run = c.paragraphs[0].runs[0]
        run.bold = True
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), "1a237e")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:val"), "clear")
        tcPr = c._tc.get_or_add_tcPr()
        tcPr.append(shd)

    for r in rows:
        vals = [r.get("host",""), r.get("service",""), r.get("port",""),
                r.get("severity",""), (r.get("finding","") or "")[:80],
                r.get("cve",""), r.get("cvss",""), r.get("scanner","")]
        row_cells = tbl.add_row().cells
        for i, v in enumerate(vals):
            row_cells[i].text = v or ""
            p = row_cells[i].paragraphs[0]
            run = p.runs[0] if p.runs else p.add_run(v or "")
            run.font.size = Pt(7)

    buf = io.BytesIO()
    doc.save(buf)
    b64 = base64.b64encode(buf.getvalue()).decode()
    ts2 = _dt.datetime.now().strftime("%Y%m%d_%H%M%S")
    return (dict(content=b64, filename=f"pentest_report_{ts2}.docx",
                 type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                 base64=True),
            f"Word doc ready — {len(rows)} findings exported.")
