"""
Ægis Red Team Suite — Reporting Module
========================================
Aggregates findings from all active scan engines and exports them as:
  • PDF   — ReportLab-generated professional pentest report (PTES / NIST 800-115)
  • CSV   — flat finding table for SIEM / ticketing system import
  • DOCX  — Word document for executive / management distribution

All exports are served via dcc.Download (in-memory, no temp files left on disk).
"""

import io
import csv
import json
import datetime
import textwrap

import dash
from dash import html, dcc, Input, Output, State, callback, no_update
import dash_bootstrap_components as dbc

# ─── Colour tokens ────────────────────────────────────────────────────────────
BG     = "#0a0808"
CARD   = "#111317"
CARD2  = "#0d0e14"
BORDER = "1px solid #1e2230"
RED    = "#ff3355"
ORANGE = "#f97316"
YELLOW = "#f5e642"
CYAN   = "#00c2ff"
GREEN  = "#00e5b0"
TEXT   = "#a2a5b5"

SEV_COL = {
    "Critical": RED,
    "High":     ORANGE,
    "Medium":   YELLOW,
    "Info":     CYAN,
    "Low":      GREEN,
}

# ─────────────────────────────────────────────────────────────────────────────
# DATA AGGREGATION
# ─────────────────────────────────────────────────────────────────────────────

def _gather_findings() -> list[dict]:
    """Pull findings from all available scan engines into a flat list of dicts."""
    findings: list[dict] = []

    # ── Ægis AD Scanner (PingCastle) ─────────────────────────────────────────
    try:
        from cyber_range.services.scan_orchestrator import ScanOrchestrator
        sess = ScanOrchestrator.get_session()
        if sess and sess.ping and sess.ping.findings:
            for f in sess.ping.findings:
                findings.append({
                    "source":      "Ægis AD Scanner",
                    "severity":    getattr(f, "severity", "Info"),
                    "category":    getattr(f, "category",  "AD / Kerberos"),
                    "title":       getattr(f, "rule",      "Finding"),
                    "description": getattr(f, "description", ""),
                    "host":        getattr(sess.ping, "domain", ""),
                    "cvss":        "",
                    "remediation": getattr(f, "remediation", "Review AD security configuration."),
                    "mitre":       "",
                })
    except Exception:
        pass

    # ── Nmap ─────────────────────────────────────────────────────────────────
    try:
        from cyber_range.services.scan_orchestrator import ScanOrchestrator
        sess = ScanOrchestrator.get_session()
        if sess and sess.nmap and sess.nmap.success:
            for p in sess.nmap.open_ports:
                port_num = getattr(p, "port", 0)
                HI  = {3389, 5985, 5986, 23, 21, 2049}
                MED = {445, 139, 135}
                sev = ("High" if port_num in HI else
                       "Medium" if port_num in MED else "Info")
                findings.append({
                    "source":      "Nmap",
                    "severity":    sev,
                    "category":    "Network Exposure",
                    "title":       f"Open Port {port_num}/{getattr(p,'protocol','tcp')} — {getattr(p,'service','')}",
                    "description": (f"Port {port_num} is open and running "
                                    f"{getattr(p,'service','')} {getattr(p,'version','')}").strip(),
                    "host":        getattr(sess.nmap, "target", ""),
                    "cvss":        "",
                    "remediation": ("Restrict access via firewall rules. Disable unused services."
                                   if sev in ("High", "Medium") else "Monitor port for unexpected activity."),
                    "mitre":       "T1046" if sev != "Info" else "",
                })
    except Exception:
        pass

    # ── BloodHound attack paths ───────────────────────────────────────────────
    try:
        from cyber_range.services.scan_orchestrator import ScanOrchestrator
        sess = ScanOrchestrator.get_session()
        if sess and sess.blood and sess.blood.success:
            for ap in sess.blood.attack_paths:
                findings.append({
                    "source":      "BloodHound",
                    "severity":    getattr(ap, "risk",         "High"),
                    "category":    "Attack Path",
                    "title":       f"{getattr(ap,'source','')} → {getattr(ap,'target','')} via {getattr(ap,'relationship','')}",
                    "description": (f"BloodHound identified a privilege escalation path: "
                                    f"{getattr(ap,'source','')} can reach "
                                    f"{getattr(ap,'target','')} via "
                                    f"{getattr(ap,'relationship','')}."),
                    "host":        getattr(sess.blood, "domain", ""),
                    "cvss":        "8.8" if getattr(ap, "risk", "") == "Critical" else "7.5",
                    "remediation": (f"Remove or restrict the {getattr(ap,'relationship','')} "
                                    f"permission from {getattr(ap,'source','')}."),
                    "mitre":       "T1484",
                })
    except Exception:
        pass

    # ── Correlation engine extra findings ─────────────────────────────────────
    try:
        from cyber_range.services.risk_correlation import RiskCorrelationEngine
        corr = RiskCorrelationEngine().correlate()
        if corr and corr.findings:
            seen = {f["title"] for f in findings}
            for f in corr.findings:
                title = getattr(f, "title", getattr(f, "rule", "Finding"))
                if title not in seen:
                    findings.append({
                        "source":      getattr(f, "source", "Correlation"),
                        "severity":    getattr(f, "severity", "Medium"),
                        "category":    getattr(f, "category", "Correlated"),
                        "title":       title,
                        "description": getattr(f, "detail", getattr(f, "description", "")),
                        "host":        "",
                        "cvss":        "",
                        "remediation": "",
                        "mitre":       "",
                    })
    except Exception:
        pass

    # ── External Penetration Testing findings ─────────────────────────────────
    try:
        from cyber_range.services import ext_pentest_engine as _ext_eng
        _latest = _ext_eng.get_latest_completed()
        if _latest:
            _ext_results = _latest.get("results", {})
            _ext_target  = _latest.get("target", "")
            for _tgt_key, _tgt_phases in _ext_results.items():
                if not isinstance(_tgt_phases, dict):
                    continue
                for _phase_key, _phase_list in _tgt_phases.items():
                    if not isinstance(_phase_list, list):
                        continue
                    for _item in _phase_list:
                        _tool = _item.get("tool", "")
                        _data = _item.get("data", {})
                        if not isinstance(_data, dict):
                            continue
                        for _f in _data.get("findings", []):
                            if _tool == "NUCLEI":
                                findings.append({
                                    "source":      "Ext Pentest / NUCLEI",
                                    "severity":    _f.get("severity", "info").capitalize(),
                                    "category":    "Vulnerability",
                                    "title":       _f.get("name", _f.get("template_id", "Nuclei finding")),
                                    "description": _f.get("description", ""),
                                    "host":        _f.get("url", _f.get("host", _ext_target)),
                                    "cvss":        str(_f.get("cvss_score", "")),
                                    "remediation": "",
                                    "mitre":       "",
                                })
                            elif _tool in ("DNS", "SSL"):
                                sev = "High" if "CRITICAL" in str(_f) else "Medium"
                                findings.append({
                                    "source":      f"Ext Pentest / {_tool}",
                                    "severity":    sev,
                                    "category":    "Cryptography" if _tool == "SSL" else "DNS / Email",
                                    "title":       str(_f),
                                    "description": "",
                                    "host":        _ext_target,
                                    "cvss":        "",
                                    "remediation": "",
                                    "mitre":       "",
                                })
                        if _tool == "ZAP":
                            for _f in _data.get("findings", []):
                                findings.append({
                                    "source":      "Ext Pentest / ZAP",
                                    "severity":    _f.get("risk", "Medium"),
                                    "category":    "Web Application",
                                    "title":       _f.get("name", _f.get("alert", "ZAP Finding")),
                                    "description": _f.get("description", ""),
                                    "host":        _f.get("url", _ext_target),
                                    "cvss":        "",
                                    "remediation": _f.get("solution", ""),
                                    "mitre":       "",
                                })
                        if _tool == "SQLMAP":
                            for _f in _data.get("findings", []):
                                findings.append({
                                    "source":      "Ext Pentest / SQLMAP",
                                    "severity":    "Critical",
                                    "category":    "Injection",
                                    "title":       f"SQL Injection: {_f.get('parameter', _f.get('url',''))}",
                                    "description": f"Injectable parameter: {_f.get('parameter','')}",
                                    "host":        _f.get("url", _ext_target),
                                    "cvss":        "9.8",
                                    "remediation": "Use parameterised queries / ORM. Apply WAF.",
                                    "mitre":       "T1190",
                                })
                        if _tool == "NIKTO":
                            for _f in _data.get("findings", []):
                                findings.append({
                                    "source":      "Ext Pentest / NIKTO",
                                    "severity":    "Medium",
                                    "category":    "Web Server",
                                    "title":       _f.get("msg", _f.get("message", "Nikto finding")),
                                    "description": "",
                                    "host":        _f.get("url", _ext_target),
                                    "cvss":        "",
                                    "remediation": "",
                                    "mitre":       "",
                                })
    except Exception:
        pass

    # ── Demo fallback (no scan yet) ───────────────────────────────────────────
    if not findings:
        findings = [
            {"source": "Ægis AD Scanner", "severity": "Critical", "category": "ACEs/DACL",
             "title": "GenericAll on Domain Admins", "description": "svc_backup has GenericAll over Domain Admins group.",
             "host": "expedite.local", "cvss": "9.8", "remediation": "Remove GenericAll ACE from svc_backup.", "mitre": "T1484"},
            {"source": "Ægis AD Scanner", "severity": "High", "category": "Kerberos",
             "title": "Kerberoastable Service Account", "description": "svc_monitor has SPN and weak password.",
             "host": "expedite.local", "cvss": "8.1", "remediation": "Rotate password, use gMSA.", "mitre": "T1558.003"},
            {"source": "Nmap", "severity": "High", "category": "Network Exposure",
             "title": "Open Port 3389/tcp — ms-wbt-server", "description": "RDP is exposed to the network.",
             "host": "192.168.1.10", "cvss": "7.5", "remediation": "Restrict RDP access via firewall / VPN.", "mitre": "T1046"},
            {"source": "Nmap", "severity": "Medium", "category": "Network Exposure",
             "title": "Open Port 445/tcp — microsoft-ds", "description": "SMB port is open.",
             "host": "192.168.1.10", "cvss": "5.9", "remediation": "Disable SMBv1; restrict SMB to required hosts.", "mitre": "T1046"},
            {"source": "BloodHound", "severity": "Critical", "category": "Attack Path",
             "title": "helpdesk → EXPEDITE-DC via WriteDACL", "description": "helpdesk user can modify DC DACL.",
             "host": "expedite.local", "cvss": "9.1", "remediation": "Remove WriteDACL from helpdesk on EXPEDITE-DC.", "mitre": "T1484"},
            {"source": "BloodHound", "severity": "High", "category": "Attack Path",
             "title": "john.doe → svc_backup via Kerberoastable", "description": "Kerberoast path to backup service account.",
             "host": "expedite.local", "cvss": "8.0", "remediation": "Rotate svc_backup credentials; enforce AES-only Kerberos.", "mitre": "T1558.003"},
        ]

    # Sort: Critical -> High -> Medium -> Info / Low
    _ORDER = {"Critical": 0, "High": 1, "Medium": 2, "Info": 3, "Low": 4}
    findings.sort(key=lambda x: _ORDER.get(x.get("severity", "Info"), 5))
    return findings


def _meta() -> dict:
    """Gather scan metadata for report headers."""
    meta = {
        "domain":    "expedite.local",
        "target":    "192.168.1.10",
        "score":     "—",
        "level":     "—",
        "timestamp": datetime.datetime.now().strftime("%Y-%m-%d %H:%M"),
        "scan_id":   "—",
    }
    try:
        from cyber_range.services.scan_orchestrator import ScanOrchestrator
        sess = ScanOrchestrator.get_session()
        if sess:
            if sess.ping:
                meta["domain"] = getattr(sess.ping, "domain", meta["domain"]) or meta["domain"]
            if sess.nmap:
                meta["target"] = getattr(sess.nmap, "target", meta["target"]) or meta["target"]
            if sess.score:
                meta["score"] = str(getattr(sess.score, "total", "—"))
                meta["level"] = getattr(sess.score, "level", "—")
            if sess.result:
                meta["scan_id"] = getattr(sess.result, "scan_id", "—") or "—"
    except Exception:
        pass
    return meta


# ─────────────────────────────────────────────────────────────────────────────
# PDF EXPORT  (ReportLab)
# ─────────────────────────────────────────────────────────────────────────────

def _build_pdf(findings: list[dict], meta: dict) -> bytes:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib import colors
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import mm
    from reportlab.platypus import (SimpleDocTemplate, Paragraph, Spacer,
                                    Table, TableStyle, HRFlowable, PageBreak)
    from reportlab.lib.enums import TA_LEFT, TA_CENTER

    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4,
                            leftMargin=18*mm, rightMargin=18*mm,
                            topMargin=18*mm, bottomMargin=18*mm)

    _BG   = colors.HexColor("#0d0e14")
    _CYAN = colors.HexColor("#00c2ff")
    _RED  = colors.HexColor("#ff3355")
    _ORG  = colors.HexColor("#f97316")
    _YLW  = colors.HexColor("#f5e642")
    _GRN  = colors.HexColor("#00e5b0")
    _TXT  = colors.HexColor("#e0e0e0")
    _CARD = colors.HexColor("#111317")
    _DARK = colors.HexColor("#080a10")
    _BRD  = colors.HexColor("#1e2230")

    _SEV_C = {"Critical": _RED, "High": _ORG, "Medium": _YLW,
              "Info": _CYAN, "Low": _GRN}

    styles = getSampleStyleSheet()
    _H1 = ParagraphStyle("h1", fontSize=22, textColor=_CYAN, leading=28,
                          spaceAfter=4, fontName="Helvetica-Bold")
    _H2 = ParagraphStyle("h2", fontSize=13, textColor=_CYAN, leading=18,
                          spaceAfter=3, fontName="Helvetica-Bold")
    _BODY = ParagraphStyle("body", fontSize=9, textColor=_TXT, leading=14,
                            spaceAfter=2, fontName="Helvetica")
    _MONO = ParagraphStyle("mono", fontSize=8, textColor=_CYAN, leading=12,
                            fontName="Courier")
    _LABEL = ParagraphStyle("label", fontSize=8, textColor=colors.HexColor("#555"),
                             leading=10, fontName="Helvetica")

    sev_counts = {}
    for f in findings:
        sev_counts[f["severity"]] = sev_counts.get(f["severity"], 0) + 1

    story = []

    # ── Cover page ────────────────────────────────────────────────────────────
    story.append(Spacer(1, 30*mm))
    story.append(Paragraph("🛡️  Ægis Red Team Suite", _H1))
    story.append(Paragraph("<b>Security Assessment Report</b>", _H2))
    story.append(HRFlowable(width="100%", thickness=1, color=_RED, spaceAfter=8))
    story.append(Spacer(1, 4*mm))

    meta_data = [
        ["Target Domain",  meta["domain"]],
        ["Primary Host",   meta["target"]],
        ["Risk Score",     f"{meta['score']} / 100  (Level {meta['level']})"],
        ["Scan Reference", meta["scan_id"]],
        ["Generated",      meta["timestamp"]],
        ["Standards",      "PTES · NIST SP 800-115 · CVSS v3.1"],
    ]
    mt = Table(meta_data, colWidths=[45*mm, 120*mm])
    mt.setStyle(TableStyle([
        ("BACKGROUND", (0,0), (-1,-1), _CARD),
        ("TEXTCOLOR",  (0,0), (0,-1), _CYAN),
        ("TEXTCOLOR",  (1,0), (1,-1), _TXT),
        ("FONTNAME",   (0,0), (0,-1), "Helvetica-Bold"),
        ("FONTSIZE",   (0,0), (-1,-1), 9),
        ("ROWBACKGROUNDS", (0,0), (-1,-1), [_CARD, _DARK]),
        ("GRID",       (0,0), (-1,-1), 0.4, _BRD),
        ("LEFTPADDING",(0,0), (-1,-1), 8),
        ("TOPPADDING", (0,0), (-1,-1), 5),
        ("BOTTOMPADDING",(0,0),(-1,-1),5),
    ]))
    story.append(mt)
    story.append(Spacer(1, 10*mm))

    # ── Severity summary bar ──────────────────────────────────────────────────
    story.append(Paragraph("Finding Summary", _H2))
    sev_row = [[s, str(sev_counts.get(s, 0))] for s in ["Critical","High","Medium","Info","Low"]]
    st = Table([[s for s,_ in sev_row], [c for _,c in sev_row]],
               colWidths=[33*mm]*5)
    st.setStyle(TableStyle([
        ("BACKGROUND", (0,0),(0,0), _RED),
        ("BACKGROUND", (1,0),(1,0), _ORG),
        ("BACKGROUND", (2,0),(2,0), _YLW),
        ("BACKGROUND", (3,0),(3,0), _CYAN),
        ("BACKGROUND", (4,0),(4,0), _GRN),
        ("TEXTCOLOR",  (0,0),(-1,-1), colors.black),
        ("FONTNAME",   (0,0),(-1,-1), "Helvetica-Bold"),
        ("FONTSIZE",   (0,0),(-1,-1), 10),
        ("ALIGN",      (0,0),(-1,-1), "CENTER"),
        ("TOPPADDING", (0,0),(-1,-1), 6),
        ("BOTTOMPADDING",(0,0),(-1,-1),6),
        ("ROWBACKGROUNDS",(0,1),(-1,-1),[_CARD, _DARK]),
    ]))
    story.append(st)
    story.append(PageBreak())

    # ── Findings detail ───────────────────────────────────────────────────────
    story.append(Paragraph("Detailed Findings", _H1))
    story.append(HRFlowable(width="100%", thickness=1, color=_RED, spaceAfter=6))

    for i, f in enumerate(findings, 1):
        sev = f["severity"]
        sev_c = _SEV_C.get(sev, _TXT)
        story.append(Spacer(1, 4*mm))
        story.append(Paragraph(f"Finding {i:02d} — {f['title']}", _H2))

        detail_data = [
            ["Severity",    sev],
            ["Source",      f["source"]],
            ["Category",    f["category"]],
            ["Host",        f["host"] or "—"],
            ["CVSS",        f["cvss"] or "—"],
            ["MITRE ATT&CK",f["mitre"] or "—"],
        ]
        dt = Table(detail_data, colWidths=[35*mm, 130*mm])
        dt.setStyle(TableStyle([
            ("BACKGROUND", (0,0), (-1,-1), _DARK),
            ("TEXTCOLOR",  (0,0), (0,-1), _CYAN),
            ("TEXTCOLOR",  (1,0), (1,-1), _TXT),
            ("FONTNAME",   (0,0), (0,-1), "Helvetica-Bold"),
            ("FONTSIZE",   (0,0), (-1,-1), 8),
            ("GRID",       (0,0), (-1,-1), 0.3, _BRD),
            ("LEFTPADDING",(0,0),(-1,-1), 6),
            ("TOPPADDING", (0,0),(-1,-1), 4),
            ("BOTTOMPADDING",(0,0),(-1,-1),4),
            # Severity cell colour
            ("BACKGROUND", (1,0), (1,0), sev_c),
            ("TEXTCOLOR",  (1,0), (1,0), colors.black),
            ("FONTNAME",   (1,0), (1,0), "Helvetica-Bold"),
        ]))
        story.append(dt)
        story.append(Spacer(1, 2*mm))

        desc = textwrap.fill(f["description"] or "No description.", 120)
        story.append(Paragraph("<b>Description:</b>", _BODY))
        story.append(Paragraph(desc, _BODY))

        if f.get("remediation"):
            rem = textwrap.fill(f["remediation"], 120)
            story.append(Paragraph("<b>Remediation:</b>", _BODY))
            story.append(Paragraph(rem, _BODY))

        story.append(HRFlowable(width="100%", thickness=0.3, color=_BRD, spaceAfter=2))

    doc.build(story)
    return buf.getvalue()


# ─────────────────────────────────────────────────────────────────────────────
# CSV EXPORT
# ─────────────────────────────────────────────────────────────────────────────

def _build_csv(findings: list[dict], meta: dict) -> str:
    buf = io.StringIO()
    writer = csv.DictWriter(buf, fieldnames=[
        "source","severity","category","title","description",
        "host","cvss","mitre","remediation"
    ])
    writer.writeheader()
    for f in findings:
        writer.writerow({k: f.get(k,"") for k in writer.fieldnames})
    return buf.getvalue()


# ─────────────────────────────────────────────────────────────────────────────
# DOCX EXPORT  (python-docx)
# ─────────────────────────────────────────────────────────────────────────────

def _build_docx(findings: list[dict], meta: dict) -> bytes:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    def _hex(h):
        h = h.lstrip("#")
        return RGBColor(int(h[:2],16), int(h[2:4],16), int(h[4:],16))

    doc = Document()

    # ── Cover ─────────────────────────────────────────────────────────────────
    t = doc.add_heading("Ægis Red Team Suite", 0)
    t.runs[0].font.color.rgb = _hex("#00c2ff")

    doc.add_heading("Security Assessment Report", 1)
    doc.add_paragraph(f"Target Domain   : {meta['domain']}")
    doc.add_paragraph(f"Primary Host    : {meta['target']}")
    doc.add_paragraph(f"Risk Score      : {meta['score']} / 100  (Level {meta['level']})")
    doc.add_paragraph(f"Scan Reference  : {meta['scan_id']}")
    doc.add_paragraph(f"Generated       : {meta['timestamp']}")
    doc.add_paragraph(f"Standards       : PTES · NIST SP 800-115 · CVSS v3.1")
    doc.add_page_break()

    # ── Summary table ─────────────────────────────────────────────────────────
    doc.add_heading("Finding Summary", 1)
    sev_counts = {}
    for f in findings:
        sev_counts[f["severity"]] = sev_counts.get(f["severity"], 0) + 1

    tbl = doc.add_table(rows=2, cols=5)
    tbl.style = "Table Grid"
    hdr_cells = tbl.rows[0].cells
    val_cells  = tbl.rows[1].cells
    for i, sev in enumerate(["Critical","High","Medium","Info","Low"]):
        hdr_cells[i].text = sev
        val_cells[i].text  = str(sev_counts.get(sev, 0))
        run = hdr_cells[i].paragraphs[0].runs[0]
        run.bold = True
        run.font.size = Pt(10)
    doc.add_page_break()

    # ── Findings ──────────────────────────────────────────────────────────────
    doc.add_heading("Detailed Findings", 1)

    for i, f in enumerate(findings, 1):
        doc.add_heading(f"Finding {i:02d} — {f['title']}", 2)
        info = [
            ("Severity",     f["severity"]),
            ("Source",       f["source"]),
            ("Category",     f["category"]),
            ("Host / Target",f["host"] or "—"),
            ("CVSS Score",   f["cvss"] or "—"),
            ("MITRE ATT&CK", f["mitre"] or "—"),
        ]
        tbl = doc.add_table(rows=len(info), cols=2)
        tbl.style = "Table Grid"
        for r, (k, v) in enumerate(info):
            tbl.rows[r].cells[0].text = k
            tbl.rows[r].cells[0].paragraphs[0].runs[0].bold = True
            tbl.rows[r].cells[1].text = v

        if f.get("description"):
            doc.add_paragraph("Description:", style="Intense Quote")
            doc.add_paragraph(f["description"])
        if f.get("remediation"):
            doc.add_paragraph("Remediation:", style="Intense Quote")
            doc.add_paragraph(f["remediation"])
        doc.add_paragraph("")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ─────────────────────────────────────────────────────────────────────────────
# DASH LAYOUT
# ─────────────────────────────────────────────────────────────────────────────

def _sev_badge(sev: str) -> html.Span:
    return html.Span(sev, style={
        "background":    SEV_COL.get(sev, "#555"),
        "color":         "#000" if sev in ("Medium","Info","Low") else "#fff",
        "padding":       "2px 8px",
        "borderRadius":  "4px",
        "fontSize":      "9px",
        "fontWeight":    "800",
        "letterSpacing": "0.5px",
    })


def _finding_row(i: int, f: dict) -> html.Tr:
    return html.Tr([
        html.Td(str(i), style={"color": "#555", "fontSize": "10px",
                               "padding": "6px 8px", "fontFamily": "monospace"}),
        html.Td(_sev_badge(f["severity"]), style={"padding": "6px 8px"}),
        html.Td(html.Span(f["source"], style={"color": CYAN, "fontSize": "10px",
                                               "fontFamily": "monospace"}),
                style={"padding": "6px 8px"}),
        html.Td(f["category"], style={"color": TEXT, "fontSize": "10px",
                                       "padding": "6px 8px"}),
        html.Td(f["title"], style={"color": "#e0e0e0", "fontSize": "10px",
                                    "fontWeight": "600", "padding": "6px 8px"}),
        html.Td(f["host"] or "—", style={"color": "#555", "fontSize": "9px",
                                          "fontFamily": "monospace", "padding": "6px 8px"}),
        html.Td(f["cvss"] or "—", style={"color": ORANGE, "fontSize": "9px",
                                           "fontWeight": "700", "padding": "6px 8px"}),
        html.Td(f["mitre"] or "—", style={"color": "#9b59b6", "fontSize": "9px",
                                            "fontFamily": "monospace", "padding": "6px 8px"}),
    ], style={"borderBottom": "1px solid #15182a",
              "transition": "background 0.15s ease"})


def layout() -> html.Div:
    """Full reporting page layout — call once per tab render."""
    findings = _gather_findings()
    meta     = _meta()

    sev_counts = {}
    for f in findings:
        sev_counts[f["severity"]] = sev_counts.get(f["severity"], 0) + 1

    src_counts = {}
    for f in findings:
        src_counts[f["source"]] = src_counts.get(f["source"], 0) + 1

    now = datetime.datetime.now().strftime("%Y-%m-%d %H:%M")

    # KPI bar
    kpis = [
        ("Total",    str(len(findings)),              "#e0e0e0"),
        ("Critical", str(sev_counts.get("Critical", 0)), RED),
        ("High",     str(sev_counts.get("High",     0)), ORANGE),
        ("Medium",   str(sev_counts.get("Medium",   0)), YELLOW),
        ("Info",     str(sev_counts.get("Info",     0)), CYAN),
    ]

    kpi_cards = []
    for label, val, col in kpis:
        kpi_cards.append(html.Div([
            html.Div(val,   style={"color": col, "fontSize": "22px",
                                   "fontWeight": "900", "fontFamily": "monospace"}),
            html.Div(label, style={"color": TEXT, "fontSize": "9px",
                                   "fontWeight": "700", "letterSpacing": "0.8px",
                                   "textTransform": "uppercase"}),
        ], style={
            "background": CARD2, "border": f"1px solid {col}33",
            "borderRadius": "8px", "padding": "12px 18px",
            "textAlign": "center", "minWidth": "80px",
        }))

    # Finding rows
    rows = [_finding_row(i+1, f) for i, f in enumerate(findings)]

    # Source badges
    src_badges = [
        html.Span(f"{src}  {cnt}", style={
            "background": {"Ægis AD Scanner": "#9b59b6",
                           "Nmap":            "#3498db",
                           "BloodHound":      "#e74c3c",
                           "Correlation":     "#e67e22"}.get(src, "#555"),
            "color": "#fff", "padding": "3px 10px", "borderRadius": "4px",
            "fontSize": "10px", "fontWeight": "700", "marginRight": "8px",
        }) for src, cnt in src_counts.items()
    ]

    BTN_S = {
        "border": "none", "borderRadius": "8px",
        "fontWeight": "700", "fontSize": "12px",
        "padding": "10px 22px", "cursor": "pointer",
        "display": "flex", "alignItems": "center", "gap": "8px",
        "transition": "all 0.2s ease",
    }

    return html.Div([
        # ── Findings snapshot — populated at page-load, consumed by callbacks ─
        dcc.Store(id="rt-rpt-findings-store", data=findings),
        dcc.Store(id="rt-rpt-meta-store",     data=meta),
        # ── Download components ─────────────────────────────────────────────
        dcc.Download(id="rt-rpt-dl-pdf"),
        dcc.Download(id="rt-rpt-dl-csv"),
        dcc.Download(id="rt-rpt-dl-docx"),

        # ── Header ──────────────────────────────────────────────────────────
        html.Div([
            html.Div([
                html.Span("📋", style={"fontSize": "18px"}),
                html.Div([
                    html.Span("ÆGIS  REPORTING", style={
                        "color": CYAN, "fontWeight": "900", "fontSize": "14px",
                        "fontFamily": "'JetBrains Mono', monospace",
                        "letterSpacing": "3px",
                    }),
                    html.Div("All findings · PTES · NIST SP 800-115 · CVSS v3.1", style={
                        "color": "#3a5a70", "fontSize": "10px", "letterSpacing": "0.5px",
                    }),
                ]),
            ], style={"display": "flex", "alignItems": "center", "gap": "12px"}),
            html.Div([
                html.Span(f"🕐 {now}", style={"color": "#333", "fontSize": "10px",
                                               "fontFamily": "monospace"}),
                html.Span(f"  {meta['domain']}", style={"color": "#3a5a70",
                                                          "fontSize": "10px",
                                                          "marginLeft": "12px"}),
            ]),
        ], style={
            "background": "linear-gradient(90deg, #060e14, #080f18)",
            "borderBottom": f"2px solid {CYAN}33",
            "padding": "14px 20px",
            "display": "flex", "justifyContent": "space-between",
            "alignItems": "center",
        }),

        # ── KPI row ─────────────────────────────────────────────────────────
        html.Div([
            *kpi_cards,
            html.Div(style={"flex": "1"}),
            html.Div(src_badges, style={"display": "flex", "alignItems": "center"}),
        ], style={
            "display": "flex", "alignItems": "center",
            "gap": "10px", "padding": "14px 20px",
            "background": CARD, "borderBottom": BORDER,
            "flexWrap": "wrap",
        }),

        # ── Export buttons ───────────────────────────────────────────────────
        html.Div([
            html.Div("📤  Export Report", style={
                "color": CYAN, "fontWeight": "700", "fontSize": "11px",
                "letterSpacing": "1px", "textTransform": "uppercase",
                "marginRight": "16px",
            }),
            html.Button([
                html.Span("📄"), "  Export PDF",
            ], id="rt-rpt-btn-pdf", n_clicks=0, style={
                **BTN_S,
                "background": "linear-gradient(135deg, #c0392b, #e74c3c)",
                "color": "#fff",
                "boxShadow": "0 4px 15px rgba(192,57,43,0.4)",
            }),
            html.Button([
                html.Span("📊"), "  Export CSV",
            ], id="rt-rpt-btn-csv", n_clicks=0, style={
                **BTN_S,
                "background": "linear-gradient(135deg, #1a6b2e, #27ae60)",
                "color": "#fff",
                "boxShadow": "0 4px 15px rgba(39,174,96,0.4)",
            }),
            html.Button([
                html.Span("📝"), "  Export Word",
            ], id="rt-rpt-btn-docx", n_clicks=0, style={
                **BTN_S,
                "background": "linear-gradient(135deg, #1a3a8a, #2980b9)",
                "color": "#fff",
                "boxShadow": "0 4px 15px rgba(41,128,185,0.4)",
            }),
            html.Div(id="rt-rpt-status", style={
                "color": GREEN, "fontSize": "11px",
                "fontFamily": "monospace", "marginLeft": "14px",
                "animation": "rt-pulse 2s ease-in-out infinite",
            }),
        ], style={
            "display": "flex", "alignItems": "center",
            "gap": "12px", "padding": "14px 20px",
            "background": "linear-gradient(90deg,#08100f,#060e14)",
            "borderBottom": BORDER, "flexWrap": "wrap",
        }),

        # ── Findings table ───────────────────────────────────────────────────
        html.Div([
            html.Div([
                html.Div("⚠️  All Findings", style={
                    "color": ORANGE, "fontWeight": "700",
                    "fontSize": "11px", "letterSpacing": "0.5px",
                }),
                html.Div(f"{len(findings)} total  ·  sorted by severity", style={
                    "color": "#333", "fontSize": "10px",
                }),
            ], style={
                "display": "flex", "justifyContent": "space-between",
                "alignItems": "center",
                "padding": "10px 16px", "borderBottom": BORDER,
                "background": CARD2,
            }),
            html.Div([
                html.Table([
                    html.Thead(html.Tr([
                        html.Th(h, style={
                            "color": CYAN, "fontSize": "9px",
                            "fontWeight": "700", "textTransform": "uppercase",
                            "letterSpacing": "0.8px", "padding": "8px 8px",
                            "borderBottom": f"2px solid {CYAN}33",
                            "whiteSpace": "nowrap",
                            "background": CARD2,
                        })
                        for h in ["#", "Severity", "Source", "Category",
                                  "Finding", "Host", "CVSS", "MITRE"]
                    ])),
                    html.Tbody(rows or [
                        html.Tr(html.Td(
                            "No findings yet — run a scan to populate this report.",
                            colSpan=8,
                            style={"color": "#333", "fontStyle": "italic",
                                   "textAlign": "center", "padding": "40px"},
                        ))
                    ]),
                ], style={
                    "width": "100%", "borderCollapse": "collapse",
                    "fontSize": "11px",
                }),
            ], style={
                "overflowX": "auto",
                "maxHeight": "calc(100vh - 300px)",
                "overflowY": "auto",
            }),
        ], style={
            "background": CARD, "border": BORDER,
            "borderRadius": "10px", "margin": "16px",
        }),

    ], style={"background": "#08090e", "minHeight": "100vh"})


# ─────────────────────────────────────────────────────────────────────────────
# CALLBACKS
# ─────────────────────────────────────────────────────────────────────────────

def register_callbacks(app):
    """Register the three download callbacks — call once from the host app.

    IMPORTANT: All three callbacks read findings from dcc.Store (rt-rpt-findings-store)
    which is populated at page-load time by layout(). This guarantees the exported
    file contains exactly the same findings that the user sees in the table — regardless
    of whether the scan session is accessible in the callback worker context.
    """

    @app.callback(
        Output("rt-rpt-dl-pdf",          "data"),
        Output("rt-rpt-status",          "children"),
        Input("rt-rpt-btn-pdf",          "n_clicks"),
        State("rt-rpt-findings-store",   "data"),
        State("rt-rpt-meta-store",       "data"),
        prevent_initial_call=True,
    )
    def _dl_pdf(n, stored_findings, stored_meta):
        if not n:
            return no_update, no_update
        # Use the exact snapshot shown in the table; fall back to live fetch only
        # if the store is somehow empty (e.g. first load race condition).
        findings = stored_findings or _gather_findings()
        meta     = stored_meta     or _meta()
        try:
            data = _build_pdf(findings, meta)
            ts   = datetime.datetime.now().strftime("%Y%m%d_%H%M")
            return (
                dcc.send_bytes(data, f"aegis_report_{ts}.pdf"),
                f"✅ PDF ready — {len(findings)} findings exported",
            )
        except Exception as e:
            return no_update, f"❌ PDF error: {e}"

    @app.callback(
        Output("rt-rpt-dl-csv",          "data"),
        Output("rt-rpt-status",          "children", allow_duplicate=True),
        Input("rt-rpt-btn-csv",          "n_clicks"),
        State("rt-rpt-findings-store",   "data"),
        State("rt-rpt-meta-store",       "data"),
        prevent_initial_call=True,
    )
    def _dl_csv(n, stored_findings, stored_meta):
        if not n:
            return no_update, no_update
        findings = stored_findings or _gather_findings()
        meta     = stored_meta     or _meta()
        try:
            csv_str = _build_csv(findings, meta)
            ts      = datetime.datetime.now().strftime("%Y%m%d_%H%M")
            return (
                dcc.send_string(csv_str, f"aegis_findings_{ts}.csv"),
                f"✅ CSV ready — {len(findings)} rows exported",
            )
        except Exception as e:
            return no_update, f"❌ CSV error: {e}"

    @app.callback(
        Output("rt-rpt-dl-docx",         "data"),
        Output("rt-rpt-status",          "children", allow_duplicate=True),
        Input("rt-rpt-btn-docx",         "n_clicks"),
        State("rt-rpt-findings-store",   "data"),
        State("rt-rpt-meta-store",       "data"),
        prevent_initial_call=True,
    )
    def _dl_docx(n, stored_findings, stored_meta):
        if not n:
            return no_update, no_update
        findings = stored_findings or _gather_findings()
        meta     = stored_meta     or _meta()
        try:
            data = _build_docx(findings, meta)
            ts   = datetime.datetime.now().strftime("%Y%m%d_%H%M")
            return (
                dcc.send_bytes(data, f"aegis_report_{ts}.docx"),
                f"✅ Word doc ready — {len(findings)} findings exported",
            )
        except Exception as e:
            return no_update, f"❌ DOCX error: {e}"
